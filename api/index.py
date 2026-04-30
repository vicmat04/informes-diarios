import os
import sys
import json
import io
import base64
import tempfile
from datetime import date, datetime
from pathlib import Path

import unicodedata
from flask import Flask, Response, jsonify, render_template, request
from fpdf import FPDF
from docx import Document
from docx.shared import Inches

# Configuración de rutas para Vercel
# APP_DIR es 'api/'
APP_DIR = Path(__file__).resolve().parent
ROOT_DIR = APP_DIR.parent
DATA_FILE = ROOT_DIR / "actividades.json"

app = Flask(__name__, template_folder='../templates', static_folder='../static')

DEFAULT_ACTIVITIES = [
    {"title": "Revisión de correos", "detail": "Revisión de correos y mensajes de WhatsApp de dinamizadores y personal de infoplaza."},
    {"title": "Atención a usuarios", "detail": "Atención a usuarios y soporte en sitio."},
    {"title": "Seguimiento de pendientes", "detail": "Seguimiento de tareas pendientes y proyectos en desarrollo."},
    {"title": "Actualizacion de sistema", "detail": "Actualización de sistemas operativos y software en equipos."},
    {"title": "Reunion de equipo", "detail": "Reunión de equipo para coordinación de actividades semanales."},
    {"title": "Soporte tecnico", "detail": "Soporte técnico a infoplazas (impresora, office, OneDrive, etc)."},
]

DEFAULT_WORKSHOP_TOPICS = [
    "Comunicacion asertiva",
    "Trabajo en equipo",
    "Liderazgo",
    "Resolucion de conflictos",
    "Planificacion",
    "Herramientas digitales",
]

FACILITATOR_REGIONAL_MAP = {
    "Licdo. Juan Quiel": "Regional de Panamá",
    "Licdo. Carlos Batista": "Regional de Chiriquí",
    "Licdo. Fernando Barria": "Regional de Veraguas",
    "Licdo. Víctor Domínguez": "Regional de Los Santos",
}

FACILITATORS = list(FACILITATOR_REGIONAL_MAP.keys())

REPORT_HEADER_LINES_TEMPLATE = [
    "Infoplazas AIP",
    "{{ regional }}",
    "*****************************",
    "Informe Diario",
    "{{ facilitator }}",
    "***Facilitador Interno***",
]

def load_data() -> tuple[list[dict], list[str]]:
    if not DATA_FILE.exists():
        return DEFAULT_ACTIVITIES.copy(), DEFAULT_WORKSHOP_TOPICS.copy()

    try:
        data = json.loads(DATA_FILE.read_text(encoding="utf-8"))
    except (json.JSONDecodeError, OSError):
        return DEFAULT_ACTIVITIES.copy(), DEFAULT_WORKSHOP_TOPICS.copy()

    raw_activities = data.get("activities", [])
    activities = []
    for item in raw_activities:
        if isinstance(item, dict) and "title" in item and "detail" in item:
            activities.append({"title": str(item["title"]).strip(), "detail": str(item["detail"]).strip()})
        elif isinstance(item, str) and item.strip():
            text = item.strip()
            title = text[:30] + "..." if len(text) > 30 else text
            activities.append({"title": title, "detail": text})

    workshop_topics = [
        str(item).strip() for item in data.get("workshop_topics", []) if str(item).strip()
    ]
    return activities or DEFAULT_ACTIVITIES.copy(), workshop_topics or DEFAULT_WORKSHOP_TOPICS.copy()

def save_data(activities: list[dict], workshop_topics: list[str]) -> None:
    payload = {"activities": activities, "workshop_topics": workshop_topics}
    try:
        DATA_FILE.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    except:
        pass

def format_display_date(iso_date: str) -> str:
    try:
        return datetime.strptime(iso_date, "%Y-%m-%d").strftime("%d/%m/%Y")
    except ValueError:
        return iso_date

def build_report_lines(report_date: str, report_items: list[str], facilitator: str = "Licdo. Juan Quiel", exit_time: str = "5:00 p.m.") -> list[str]:
    regional = FACILITATOR_REGIONAL_MAP.get(facilitator, "Regional de Panamá")
    lines = []
    for line in REPORT_HEADER_LINES_TEMPLATE:
        if "{{ facilitator }}" in line:
            lines.append(line.replace("{{ facilitator }}", facilitator))
        elif "{{ regional }}" in line:
            lines.append(line.replace("{{ regional }}", regional))
        else:
            lines.append(line)
            
    lines.append(f"Fecha: {format_display_date(report_date)}")
    lines.append("")
    lines.append("Actividades realizadas:")
    
    final_items = ["Se realiza marcacion de entrada a las 8:00 a.m."]
    if report_items:
        final_items.extend(report_items)
    
    final_items.append("Se realizó la confección y envío de los informes diarios de actividades a la Licenciada Ilsa, manteniendo en copia a la Licenciada Madai para su debido seguimiento y control.")
    final_items.append(f"Se realiza marcacion de salida a las {exit_time}")

    for index, item in enumerate(final_items, start=1):
        lines.append(f"{index}. {item}")

    return lines

def build_doc_bytes(report_date: str, report_items: list[str], facilitator: str, exit_time: str, images: list[str] = None) -> bytes:
    doc = Document()
    
    logo_path = ROOT_DIR / "logo.png" if (ROOT_DIR / "logo.png").exists() else (ROOT_DIR / "logo.jpg" if (ROOT_DIR / "logo.jpg").exists() else None)
    if logo_path:
        try:
            p = doc.add_paragraph()
            p.alignment = 2
            r = p.add_run()
            r.add_picture(str(logo_path), width=Inches(1.5))
        except Exception as e:
            print("Error adding logo to docx:", e)

    lines = build_report_lines(report_date, report_items, facilitator, exit_time)
    for line in lines:
        doc.add_paragraph(line)
    
    if images:
        doc.add_page_break()
        doc.add_heading("Anexos", level=1)
        for img_b64 in images:
            if "," in img_b64:
                img_b64 = img_b64.split(",")[1]
            img_data = base64.b64decode(img_b64)
            doc.add_picture(io.BytesIO(img_data), width=Inches(6.0))
            
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

def build_pdf_bytes(report_date: str, report_items: list[str], facilitator: str, exit_time: str, images: list[str] = None) -> bytes:
    lines = build_report_lines(report_date, report_items, facilitator, exit_time)
    
    pdf = FPDF()
    pdf.set_margins(left=20, top=20, right=20)
    pdf.add_page()
    
    logo_path = ROOT_DIR / "logo.png" if (ROOT_DIR / "logo.png").exists() else (ROOT_DIR / "logo.jpg" if (ROOT_DIR / "logo.jpg").exists() else None)
    if logo_path:
        try:
            pdf.image(str(logo_path), x=15, y=15, w=40)
        except Exception as e:
            print("Error adding logo to pdf:", e)
    
    for index, line in enumerate(lines):
        if not line.strip():
            pdf.ln(5)
            continue
            
        safe_line = line.replace("“", '"').replace("”", '"').replace("‘", "'").replace("’", "'").replace("–", "-").replace("—", "-").replace("\u200b", "")
        safe_line = safe_line.encode('latin-1', 'replace').decode('latin-1')
        
        if index <= 5:
            pdf.set_font("helvetica", style="B", size=12)
            pdf.multi_cell(0, 7, txt=safe_line, align='C', new_x="LMARGIN", new_y="NEXT")
        elif "Fecha:" in safe_line or "Actividades realizadas:" in safe_line:
            pdf.set_font("helvetica", style="B", size=12)
            pdf.multi_cell(0, 7, txt=safe_line, align='L', new_x="LMARGIN", new_y="NEXT")
        else:
            pdf.set_font("helvetica", style="", size=12)
            pdf.multi_cell(0, 7, txt=safe_line, align='L', new_x="LMARGIN", new_y="NEXT")

    if images:
        pdf.add_page()
        pdf.set_font("helvetica", "B", 16)
        pdf.cell(0, 10, "Anexos", align="C", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(5)
        for img_b64 in images:
            try:
                if "," in img_b64:
                    img_b64 = img_b64.split(",")[1]
                img_data = base64.b64decode(img_b64)
                with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp:
                    tmp.write(img_data)
                    tmp_path = tmp.name
                
                pdf.image(tmp_path, w=170)
                pdf.ln(10)
                os.unlink(tmp_path)
            except Exception as e:
                print("Error adding image to pdf:", e)
        
    return bytes(pdf.output())

@app.get("/")
def index():
    activities, workshop_topics = load_data()
    
    # Mapa de avatares hardcodeado para Vercel
    # Esto asegura que las imágenes se carguen correctamente sin depender del escaneo del sistema de archivos
    avatar_map = {
        "juan quiel": "juan_quiel.jpeg",
        "carlos batista": "carlos_batista.png",
        "fernando barria": "fernando_barria.png",
        "victor dominguez": "victor_dominguez.jpeg"
    }

    initial_state = {
        "today": date.today().isoformat(),
        "activities": activities,
        "workshopTopics": workshop_topics,
        "facilitators": FACILITATORS,
        "regionalMap": FACILITATOR_REGIONAL_MAP,
        "avatarMap": avatar_map,
    }
    return render_template("index.html", initial_state=initial_state)

@app.get("/api/data")
def get_data():
    activities, workshop_topics = load_data()
    return jsonify({"activities": activities, "workshop_topics": workshop_topics})

@app.post("/api/data")
def update_data():
    payload = request.get_json(silent=True) or {}
    activities = payload.get("activities", [])
    workshop_topics = payload.get("workshop_topics", [])
    save_data(activities, workshop_topics)
    return jsonify({"ok": True})

@app.post("/api/export/<doc_type>")
def export_report(doc_type):
    data = request.json
    report_date = data.get('report_date', date.today().isoformat())
    report_items = data.get('report_items', [])
    facilitator = data.get('facilitator', "Licdo. Juan Quiel")
    exit_time = data.get('exit_time', "5:00 p.m.")
    images = data.get('images', [])

    try:
        if doc_type == 'doc':
            out_bytes = build_doc_bytes(report_date, report_items, facilitator, exit_time, images)
            return Response(out_bytes, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document', 
                            headers={"Content-Disposition": f"attachment; filename=informe_{report_date}.docx"})
        elif doc_type == 'pdf':
            out_bytes = build_pdf_bytes(report_date, report_items, facilitator, exit_time, images)
            return Response(out_bytes, mimetype='application/pdf', 
                            headers={"Content-Disposition": f"attachment; filename=informe_{report_date}.pdf"})
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    return jsonify({"error": "Invalid format"}), 400

if __name__ == "__main__":
    app.run(debug=True, port=5000)
